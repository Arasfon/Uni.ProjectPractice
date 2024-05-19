import docx
import docx.shared
import openpyxl

from math import ceil

from functions import calculate


def tabulate(x_min, x_max, x_step, y=1, a=1, b=1):
    document = docx.Document()

    document.add_paragraph("Задание")
    document.add_paragraph("Написать программу, табулирующую кусочно-заданную функцию.")
    document.add_paragraph(
        "Для выполнения задания требуется реализовать два модуля: один для табулирования, "
        + "другой для кусочно-заданной функции. В модуле кусочно-заданной функции должна быть "
        + "функция calculate, которая должна вычислять значение кусочно-заданной функции."
    )
    document.add_paragraph(
        "В основной части программы следует реализовать создание документа Microsoft Word "
        + "и рабочей книги Microsoft Excel. "
        + "В созданный документ программа должна записывать задание, "
        + "результаты работы программы (JPEG-изображение), "
        + "а также таблицу с результатами табулирования функции. "
        + "В рабочую книгу требуется записать результаты табулирования функции."
    )
    document.add_paragraph("Результаты табулирования также следует вывести в стандартный поток вывода.")
    document.add_paragraph(
        "В конце программы требуется считать созданную рабочую книгу и вывести значения, "
        + "записанные в её ячейках в стандартный поток вывода."
    )
    document.add_paragraph("Конечный результат следует обрезать до 3-х знаков (за исключением записи в Excel).")
    document.add_paragraph("Продемонстрировать результат выполнения программы — вывод во всех файлах и в стандартный поток вывода.")
    document.add_picture("result.jpg", width=docx.shared.Cm(10))

    document_table = document.add_table(1, 2, "Table Grid")
    document_table.rows[0].cells[0].text = "X"
    document_table.rows[0].cells[1].text = "F"

    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    worksheet.cell(1, 1).value = "X"
    worksheet.cell(1, 2).value = "F"

    count = ceil((x_max - x_min) / x_step)
    for i in range(count + 1):
        current_x = x_min + x_step * i
        current_f = calculate(current_x, y, a, b)

        formatted_x = str(current_x)
        formatted_f = f"{current_f:.3f}"

        print(f"x: {formatted_x}; f: {formatted_f}")

        current_document_table_row = document_table.add_row()
        current_document_table_row.cells[0].text = formatted_x
        current_document_table_row.cells[1].text = formatted_f

        worksheet.cell(i + 2, 1).value = current_x
        worksheet.cell(i + 2, 2).value = current_f

    document.save("result.docx")
    workbook.save("result.xlsx")

    workbook = openpyxl.load_workbook("result.xlsx")
    worksheet = workbook.active

    for row in worksheet.iter_rows():
        for cell in row:
            print(cell.value, end="\t")
        print()
